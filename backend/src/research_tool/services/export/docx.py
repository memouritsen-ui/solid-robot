"""DOCX export provider using python-docx."""

from datetime import datetime
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from .exporter import Exporter, ExportFormat, ExportResult, ResearchExportData


class DOCXExporter(Exporter):
    """Export research results to DOCX format using python-docx."""

    @property
    def format(self) -> ExportFormat:
        """Return the export format."""
        return ExportFormat.DOCX

    @property
    def mime_type(self) -> str:
        """Return the MIME type."""
        return (
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.document"
        )

    @property
    def file_extension(self) -> str:
        """Return the file extension."""
        return "docx"

    async def export(self, data: ResearchExportData) -> ExportResult:
        """Export research data to DOCX.

        Args:
            data: Research data to export

        Returns:
            ExportResult: DOCX content as bytes
        """
        try:
            doc = self._generate_document(data)
            buffer = BytesIO()
            doc.save(buffer)
            docx_bytes = buffer.getvalue()

            return ExportResult(
                success=True,
                format=self.format,
                content=docx_bytes,
                filename=self.generate_filename(data.query),
                mime_type=self.mime_type,
            )
        except Exception as e:
            return ExportResult(
                success=False,
                format=self.format,
                content=None,
                filename=self.generate_filename(data.query),
                mime_type=self.mime_type,
                error=str(e),
            )

    def _generate_document(self, data: ResearchExportData) -> Any:
        """Generate DOCX document from research data.

        Args:
            data: Research data

        Returns:
            Document: python-docx Document object
        """
        doc = Document()

        # Title
        title = doc.add_heading(f"Research Report: {data.query}", level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Metadata section
        doc.add_heading("Metadata", level=1)
        meta_table = doc.add_table(rows=3, cols=2)
        meta_table.style = "Table Grid"

        cells = meta_table.rows[0].cells
        cells[0].text = "Domain"
        cells[1].text = data.domain

        cells = meta_table.rows[1].cells
        cells[0].text = "Confidence Score"
        cells[1].text = f"{data.confidence_score:.1%}"

        cells = meta_table.rows[2].cells
        cells[0].text = "Generated"
        cells[1].text = datetime.now().strftime("%Y-%m-%d %H:%M")

        doc.add_paragraph()

        # Executive Summary
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(data.summary)

        # Key Findings
        if data.facts:
            doc.add_heading("Key Findings", level=1)

            for i, fact in enumerate(data.facts, 1):
                doc.add_heading(f"Finding {i}", level=2)

                statement = fact.get("statement", fact.get("content", ""))
                confidence = fact.get("confidence", 0.0)
                source = fact.get("source", "Unknown")

                # Add quote paragraph
                quote = doc.add_paragraph()
                quote.add_run(f'"{statement}"').italic = True

                # Add metadata
                meta = doc.add_paragraph()
                meta.add_run(f"Confidence: {confidence:.1%}").bold = True
                meta.add_run(f" | Source: {source}")

        # Sources
        if data.sources:
            doc.add_heading("Sources", level=1)

            for _i, source in enumerate(data.sources, 1):
                title = source.get("title", "Untitled")
                url = source.get("url", "")
                source_type = source.get("type", "web")

                p = doc.add_paragraph(style="List Number")
                p.add_run(title).bold = True
                if url:
                    p.add_run(f"\n{url}")
                p.add_run(f"\nType: {source_type}")

        # Limitations
        if data.limitations:
            doc.add_heading("Limitations", level=1)
            doc.add_paragraph(
                "This research has the following limitations that should be "
                "considered when interpreting the results:"
            )

            for limitation in data.limitations:
                doc.add_paragraph(limitation, style="List Bullet")

        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.add_run(
            "This report was generated automatically. "
            "Please verify critical information from primary sources."
        ).italic = True

        return doc
